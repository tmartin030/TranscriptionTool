# src/document/document_generator.py

from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def add_hyperlink(paragraph, run_text, bookmark):
    """Add a hyperlink to a bookmark in the given paragraph."""
    part = paragraph.part
    # Create the w:hyperlink tag and add needed attributes
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark)
    # Create a w:r element
    new_run = OxmlElement('w:r')
    # Create a w:rPr element
    new_run_rPr = OxmlElement('w:rPr')
    new_run.append(new_run_rPr)
    # Set the text for the run
    text_elem = OxmlElement('w:t')
    text_elem.text = run_text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_bookmark(paragraph, bookmark):
    """Add a bookmark to a paragraph."""
    p = paragraph._p
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:name'), bookmark)
    bookmark_start.set(qn('w:id'), '0')
    p.append(bookmark_start)
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')
    p.append(bookmark_end)

def generate_transcript_document(transcript_items):
    """
    Creates a Document containing all transcripts with a clickable Table of Contents.
    
    Args:
        transcript_items (list of dict): Each dict should have:
            - 'file_path': The original file name or identifier.
            - 'header_text': A header for the transcript.
            - 'transcription_date': A date string.
            - 'segments': A list of segments, each with keys 'start_time', 'speaker', 'transcription'.
    
    Returns:
        Document: A docx Document object with all the transcripts.
    """
    doc = Document()
    # Set default style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(12)
    
    # Title for the document
    doc.add_heading("All Transcripts", level=1)
    
    # Table of Contents Section
    doc.add_heading("Table of Contents", level=2)
    for item in transcript_items:
        header_text = item.get('header_text', 'Transcript')
        # Use a sanitized version of the header_text as the bookmark (replace spaces with underscores)
        bookmark = header_text.replace(" ", "_")
        toc_para = doc.add_paragraph()
        add_hyperlink(toc_para, header_text, bookmark)
    
    # Add each transcript
    for item in transcript_items:
        header_text = item.get('header_text', 'Transcript')
        file_path = item.get('file_path', 'Unknown File')
        transcription_date = item.get('transcription_date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        segments = item.get('segments', [])
        
        header = doc.add_heading(header_text, level=1)
        bookmark = header_text.replace(" ", "_")
        add_bookmark(header, bookmark)
        
        doc.add_paragraph(f"Audio File: {file_path}")
        doc.add_paragraph(f"Transcription Date: {transcription_date}")
        
        for seg in segments:
            start_time = seg.get('start_time', 0.0)
            speaker = seg.get('speaker', 'Unknown')
            transcription = seg.get('transcription', '')
            segment_start_time = f"{int(start_time // 3600):02d}:{int((start_time % 3600) // 60):02d}:{int(start_time % 60):02d}"
            try:
                speaker_label = int(speaker.split('_')[1]) + 1 if '_' in speaker else speaker
            except (IndexError, ValueError):
                speaker_label = speaker
            doc.add_paragraph(f"Speaker {speaker_label}: [{segment_start_time}] {transcription}")
        
        doc.add_page_break()
    
    return doc
