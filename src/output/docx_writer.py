from docx import Document
from docx.shared import Pt
from pyannote.core import Segment
from typing import List, Tuple
import datetime
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxWriter:
    def __init__(self, output_dir):
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        self.output_dir = output_dir

    def create_docx(self, audio_file: str, diarization_result: List[Tuple[Segment, str]], transcriptions: List[str], output_filename: str) -> None:
        """
        Creates a DOCX file with the transcription and diarization information.

        Args:
            audio_file: Path to the audio file.
            diarization_result: List of (Segment, speaker) tuples from the diarization.
            transcriptions: List of transcriptions.
            output_filename: Path to save the output DOCX file.
        """
        document = Document()
        document.add_heading("Transcription and Diarization", level=1)

        metadata = document.add_paragraph()
        metadata.add_run(f"Audio File: {audio_file}\n").bold = True
        metadata.add_run(f"Transcription Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

        current_speaker = None
        for (segment, speaker), transcription in zip(diarization_result, transcriptions):
            if speaker != current_speaker:
                if current_speaker is not None:
                    document.add_paragraph()

                current_speaker = speaker

                speaker_paragraph = document.add_paragraph()
                speaker_run = speaker_paragraph.add_run(f"Speaker {speaker}: ")
                speaker_run.bold = True
                speaker_run.font.size = Pt(14)

            time_str = f"[{segment.start:.2f} - {segment.end:.2f}]: "

            document.paragraphs[-1].add_run(time_str).italic = True
            document.paragraphs[-1].add_run(transcription)

        output_path = os.path.join(self.output_dir, output_filename)
        document.save(output_path)
        print(f"Transcript saved to: {output_path}")
